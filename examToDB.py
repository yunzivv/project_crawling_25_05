import re
import os
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

# ocr-env_examToDB\Scripts\activate
# python examToDB.py

# 문서 블록 순회
def iter_block_items(parent):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# 표 안 여부 확인 함수
def is_paragraph_in_table(paragraph: Paragraph):
    parent = paragraph._element
    while parent is not None:
        if parent.tag.endswith("tbl"):
            return True
        parent = parent.getparent()
    return False

# 제목 정보 추출
def extract_title_info(filename):
    basename = os.path.basename(filename)
    name, _ = os.path.splitext(basename)
    match = re.match(r'(.+?)(\d{8})$', name)
    if match:
        return match.group(1).strip(), match.group(2)
    return name, None

# 문단 앞에 새로운 문단 삽입
def insert_paragraph_before(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._element.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para

# 채워진 번호를 비채워진 번호로 변환
def replace_filled_numbers(paragraph, counter):
    filled_to_unfilled = {
        "❶": "①", "❷": "②", "❸": "③", "❹": "④",
    }
    for run in paragraph.runs:
        for filled, unfilled in filled_to_unfilled.items():
            if filled in run.text:
                run.text = run.text.replace(filled, unfilled)
                run.bold = False
                counter[0] += 1

# CBT 안내문 제거
def remove_cbt_notice(paragraphs):
    start_idx, end_idx = None, None
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if start_idx is None and text.startswith("전자문제집 CBT"):
            start_idx = i
        if start_idx is not None and text.endswith("확인하세요."):
            end_idx = i
            break
    if start_idx is not None and end_idx is not None:
        for i in range(start_idx, end_idx + 1):
            paragraphs[i]._element.getparent().remove(paragraphs[i]._element)
        print(f"🗑️ 안내문 삭제 완료: 문단 {start_idx} ~ {end_idx}")
    else:
        print("⚠️ 안내문 텍스트를 찾지 못했습니다.")

# 과목표 제거 및 텍스트만 남기기
def convert_subject_tables(doc):
    tables = list(iter_block_items(doc))
    count = 0
    for i, tbl in enumerate(tables):
        if not isinstance(tbl, Table):
            continue
        if i == len(tables) - 1:
            continue  # 마지막 정답표는 무시
        if len(tbl.rows) == 1 and len(tbl.rows[0].cells) == 1:
            cell = tbl.rows[0].cells[0]
            cell_text = cell.text.strip()
            if cell_text:
                first_para = cell.paragraphs[0]
                first_para.text = f"(Subject) {cell_text} (Subject)"
                tbl._element.getparent().insert(tbl._element.getparent().index(tbl._element), first_para._element)
                tbl._element.getparent().remove(tbl._element)
                count += 1
    print(f"🧹 과목표 변환 및 삭제 완료: {count}개")

# <<<QUESTION>>> 삽입 및 번호 변환
def insert_question_and_convert(doc):
    paragraphs = []
    for b in iter_block_items(doc):
        if isinstance(b, Paragraph):
            paragraphs.append(b)
        elif isinstance(b, Table):
            for row in b.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        paragraphs.append(para)

    print(f"\n📄 전체 문단 수: {len(paragraphs)}")
    counter = [0]
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        replace_filled_numbers(p, counter)
        bold = any(run.bold for run in p.runs if run.text.strip())
        if bold and re.match(r"^\d+\.\s", text):
            insert_paragraph_before(p, "<<<QUESTION>>>")
    print(f"✅ 숫자 변환: 총 {counter[0]}개")
    return paragraphs

# 선택지를 개별 문단으로 분리하고 마킹
def split_choice_paragraphs(doc):
    pattern = r"(①|②|③|④)"
    new_paragraphs = []

    for paragraph in list(doc.paragraphs):
        if is_paragraph_in_table(paragraph):
            continue

        full_text = paragraph.text
        if not any(opt in full_text for opt in ["①", "②", "③", "④"]):
            continue

        # ① 앞에 [choice]\n 삽입
        full_text = full_text.replace("①", "[choice]\n①")
        
        # 선택지 번호 앞에 개행 삽입
        split_text = re.sub(pattern, r"\n\1", full_text)

        # 문장 분리
        lines = [line.strip() for line in split_text.split('\n') if line.strip()]
        if len(lines) <= 1:
            continue

        # 원래 문단 삭제
        parent_elm = paragraph._element.getparent()
        insert_idx = list(parent_elm).index(paragraph._element)
        parent_elm.remove(paragraph._element)

        # 새 문단들 삽입
        for i, line in enumerate(lines):
            new_p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = line
            r.append(t)
            new_p.append(r)
            parent_elm.insert(insert_idx + i, new_p)

    print("✅ 선택지를 문단 단위로 완전히 분리 완료")

# 과목별 문제 개수 
def count_questions_in_subject(doc):
    subject_counts = {}
    current_subject = None
    current_count = 0
    subject_index = 1

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()

            # 과목 발견
            if text.startswith("(Subject)") and text.endswith("(Subject)"):
                if current_subject:
                    subject_counts[f"{subject_index}과목 : {current_subject}"] = current_count
                    subject_index += 1
                current_subject = text.replace("(Subject)", "").strip()
                current_count = 0

            # 문제 마커 발견
            elif text == "<<<QUESTION>>>":
                current_count += 1

    # 마지막 과목 저장
    if current_subject:
        subject_counts[f"{subject_index}과목 : {current_subject}"] = current_count

    print("\n📊 과목별 문제 개수:")
    for subject, count in subject_counts.items():
        print(f"  - {subject}: {count}문제")

# 문제, 선택지 분류
def parse_questions_and_choices(doc: Document):
    paragraphs = doc.paragraphs
    questions = []
    current_block = []
    
    # 블록을 문제 단위로 나누기
    for para in paragraphs:
        if para.text.strip() == "<<<QUESTION>>>":
            if current_block:
                questions.append(current_block)
            current_block = []
        else:
            current_block.append(para)
    if current_block:
        questions.append(current_block)

    parsed = []
    for block in questions:
        question_number = None
        question_text = ""
        has_image = False
        choices = []

        for para in block:
            text = para.text.strip()

            # 문제 번호 추출
            if question_number is None:
                match = re.match(r"^(\d+)\.\s*", text)
                if match:
                    question_number = int(match.group(1))

            # 이미지 포함 여부
            if not has_image:
                has_image = any("graphic" in run._element.xml for run in para.runs)

            # 선택지 추출
            if "[choice]" in text:
                # 번호 매핑
                for run in para.runs:
                    choice_text = run.text.strip()
                    if not choice_text:
                        continue
                    match = re.match(r"\[choice\]\s*(①|②|③|④)(.*)", choice_text)
                    if match:
                        choice_number = "①②③④".index(match.group(1)) + 1
                        choice_body = match.group(2).strip()
                        if choice_body:  # 텍스트 없는 경우 제외
                            choices.append((choice_number, choice_body))

            # 문제 텍스트 축적
            elif re.match(r"^\d+\.\s", text):  # 문제 본문
                question_text = text

        parsed.append({
            "number": question_number,
            "text": question_text,
            "has_image": has_image,
            "choices": choices
        })

    return parsed

# 메인 실행
def main(path):
    title, date = extract_title_info(path)
    print(f"\n📄 문서: {os.path.basename(path)}")
    doc = Document(path)

    # 안내문 삭제 먼저
    all_paragraphs = [p for b in iter_block_items(doc) if isinstance(b, Paragraph) for p in [b]]
    remove_cbt_notice(all_paragraphs)

    # 과목표 -> 텍스트로 변환 후 삭제
    convert_subject_tables(doc)

    # 문제 마킹 및 숫자 변환
    paragraphs = insert_question_and_convert(doc)

    # 선택지 문단 분리
    split_choice_paragraphs(doc)

    # 과목별 문제개수
    count_questions_in_subject(doc)

    output_path = f"{os.path.basename(path)}"
    doc.save(output_path)
    print(f"✅ 저장 완료: {output_path}")

    data = parse_questions_and_choices(doc)

    for q in data[:6]:  # 처음 5문제만 출력 예시
        print(f"\n문제 {q['number']}: {q['text']}")
        print(f"이미지 포함: {'✅' if q['has_image'] else '❌'}")
        for num, content in q['choices']:
            print(f"  {num}. {content}")

# 전체 반복
INPUT_FOLDER = "기출문제Docx"
OUTPUT_FOLDER = "기출문제포맷"

def batch_format_documents():
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)

    files = [f for f in os.listdir(INPUT_FOLDER) if f.endswith(".docx")]
    print(f"📁 총 {len(files)}개 파일 처리 시작")

    for idx, filename in enumerate(files, 1):
        input_path = os.path.join(INPUT_FOLDER, filename)
        output_path = os.path.join(OUTPUT_FOLDER, filename)
        try:
            print(f"\n[{idx}/{len(files)}] ▶ 처리 중: {filename}")
            process_single_file(input_path, output_path)
        except Exception as e:
            print(f"❌ 오류 발생 - {filename}: {e}")

def process_single_file(input_path, output_path):
    title, date = extract_title_info(input_path)
    doc = Document(input_path)

    all_paragraphs = [p for b in iter_block_items(doc) if isinstance(b, Paragraph) for p in [b]]
    remove_cbt_notice(all_paragraphs)
    convert_subject_tables(doc)
    insert_question_and_convert(doc)
    split_choice_paragraphs(doc)
    count_questions_in_subject(doc)

    doc.save(output_path)
    print(f"✅ 저장 완료: {output_path}")

if __name__ == "__main__":
    batch_format_documents()